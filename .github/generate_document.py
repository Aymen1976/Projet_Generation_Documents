import json
import os
import sys
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

def generate_docx(titre, date, contenu, output_path):
    document = Document()
    document.add_heading(titre, level=0)
    document.add_paragraph(f"Date: {date}")
    document.add_paragraph("")  # Espace vide
    document.add_paragraph(contenu)
    document.save(output_path)
    print(f"DOCX enregistré à : {output_path}")

def generate_pdf(titre, date, contenu, output_path):
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter
    textobject = c.beginText(50, height - 50)
    textobject.setFont("Helvetica", 12)
    textobject.textLine(titre)
    textobject.textLine(f"Date: {date}")
    textobject.textLine("")
    for line in contenu.splitlines():
        textobject.textLine(line)
    c.drawText(textobject)
    c.showPage()
    c.save()
    print(f"PDF enregistré à : {output_path}")

def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_document.py <fichier_json>")
        sys.exit(1)
    input_json_file = sys.argv[1]
    with open(input_json_file, 'r', encoding='utf-8-sig') as f:
        data = json.load(f)

    format_choice = data.get("format", "DOCX").upper()
    titre = data.get("titre", "Document généré par Chatbot")
    date = data.get("date", "Date inconnue")
    contenu = data.get("contenu", "")

    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    if format_choice == "PDF":
        output_file = os.path.join(desktop_path, "document_chatbot.pdf")
        generate_pdf(titre, date, contenu, output_file)
    else:
        output_file = os.path.join(desktop_path, "document_chatbot.docx")
        generate_docx(titre, date, contenu, output_file)

if __name__ == "__main__":
    main()
